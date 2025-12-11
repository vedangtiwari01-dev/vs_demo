import PyPDF2
from docx import Document
from typing import Dict, Any

class SOPExtractor:
    """Extracts text from SOP documents (PDF and DOCX)"""

    @staticmethod
    def extract_from_pdf(file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ''
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'file_type': 'pdf'
                }

                for page in pdf_reader.pages:
                    text += page.extract_text() + '\n'

                return {
                    'text': text.strip(),
                    'metadata': metadata
                }
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")

    @staticmethod
    def extract_from_docx(file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ''

            for paragraph in doc.paragraphs:
                text += paragraph.text + '\n'

            metadata = {
                'paragraphs': len(doc.paragraphs),
                'file_type': 'docx'
            }

            return {
                'text': text.strip(),
                'metadata': metadata
            }
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")

    @staticmethod
    def extract(file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract text based on file type"""
        if file_type.lower() == 'pdf':
            return SOPExtractor.extract_from_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return SOPExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
